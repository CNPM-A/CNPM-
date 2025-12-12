from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('HƯỚNG DẪN KẾT NỐI API AUTHENTICATION LOGIN', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('Smart School Bus - Frontend React/TypeScript')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(12)
subtitle_format.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()

# Section 1: Overview
doc.add_heading('I. TỔNG QUAN', 1)
doc.add_paragraph(
    'Để kết nối API Authentication Login, bạn cần thiết lập 4 file chính và '
    'hiểu rõ flow hoạt động của hệ thống xác thực.'
)

# Section 2: Required Files
doc.add_heading('II. CÁC FILE CẦN THIẾT', 1)

# File 1
doc.add_heading('1. File api.js - HTTP Client Cơ Bản', 2)
doc.add_paragraph('📁 Đường dẫn: src/services/api.js')
doc.add_paragraph('🎯 Mục đích:')
points = doc.add_paragraph()
points.style = 'List Bullet'
points.add_run('Tạo axios instance với baseURL của backend')
points = doc.add_paragraph()
points.style = 'List Bullet'
points.add_run('Tự động thêm JWT token vào headers cho mọi request')
points = doc.add_paragraph()
points.style = 'List Bullet'
points.add_run('Xử lý lỗi 401 (Unauthorized)')

doc.add_paragraph('💻 Code mẫu:')
code1 = doc.add_paragraph('''import axios from 'axios';

const API_BASE_URL = import.meta.env.VITE_API_URL || 'https://smart-school-bus-api.onrender.com/api/v1';

const api = axios.create({
  baseURL: API_BASE_URL,
  timeout: 30000,
  headers: {
    'Content-Type': 'application/json',
  },
});

// Request Interceptor: Tự động thêm token
api.interceptors.request.use(
  (config) => {
    const token = localStorage.getItem('token');
    if (token) {
      config.headers.Authorization = `Bearer ${token}`;
    }
    return config;
  },
  (error) => Promise.reject(error)
);

// Response Interceptor: Xử lý lỗi 401
api.interceptors.response.use(
  (response) => response,
  (error) => {
    if (error.response?.status === 401) {
      localStorage.removeItem('token');
      localStorage.removeItem('user');
      window.location.href = '/parent/login';
    }
    return Promise.reject(error);
  }
);

export default api;''')
code1.style = 'No Spacing'
for run in code1.runs:
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 128)

doc.add_page_break()

# File 2
doc.add_heading('2. File authService.js - Service Xử Lý Authentication', 2)
doc.add_paragraph('📁 Đường dẫn: src/services/authService.js')
doc.add_paragraph('🎯 Chức năng chính:')

funcs = doc.add_paragraph()
funcs.style = 'List Bullet'
funcs.add_run('login(credentials) - Đăng nhập')
funcs = doc.add_paragraph()
funcs.style = 'List Bullet'
funcs.add_run('logout() - Đăng xuất')
funcs = doc.add_paragraph()
funcs.style = 'List Bullet'
funcs.add_run('getToken() - Lấy token')
funcs = doc.add_paragraph()
funcs.style = 'List Bullet'
funcs.add_run('getCurrentUser() - Lấy thông tin user')
funcs = doc.add_paragraph()
funcs.style = 'List Bullet'
funcs.add_run('isAuthenticated() - Kiểm tra đã login chưa')

doc.add_paragraph('💻 Code mẫu:')
code2 = doc.add_paragraph('''import api from './api';

const login = async (credentials) => {
  try {
    // Gọi API endpoint /auth/signin
    const response = await api.post('/auth/signin', credentials);
    
    if (response.data && response.data.accessToken) {
      const { accessToken, data } = response.data;
      
      // Lưu token và user info vào localStorage
      localStorage.setItem('token', accessToken);
      if (data && data.user) {
        localStorage.setItem('user', JSON.stringify(data.user));
      }
      
      return response.data;
    } else {
      throw new Error('Login did not return an access token.');
    }
  } catch (error) {
    console.error('Login failed:', error);
    throw error;
  }
};

const logout = async () => {
  try {
    await api.delete('/auth/logout');
  } catch (error) {
    console.warn("Logout API call failed:", error);
  } finally {
    localStorage.removeItem('token');
    localStorage.removeItem('user');
  }
};

const getToken = () => {
  return localStorage.getItem('token');
};

const getCurrentUser = () => {
  const user = localStorage.getItem('user');
  try {
    return user ? JSON.parse(user) : null;
  } catch (e) {
    console.error("Failed to parse user data", e);
    return null;
  }
};

const isAuthenticated = () => {
  return !!getToken();
};

const authService = {
  login,
  logout,
  getToken,
  getCurrentUser,
  isAuthenticated,
};

export default authService;''')
code2.style = 'No Spacing'
for run in code2.runs:
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 128)

doc.add_page_break()

# File 3
doc.add_heading('3. File authService.d.ts - TypeScript Type Definitions', 2)
doc.add_paragraph('📁 Đường dẫn: src/services/authService.d.ts')
doc.add_paragraph('🎯 Mục đích: Định nghĩa types cho TypeScript')

doc.add_paragraph('💻 Code mẫu:')
code3 = doc.add_paragraph('''export interface User {
    _id: string;
    name: string;
    email: string;
    phoneNumber: string;
    role: string;
    [key: string]: any;
}

export interface AuthResponse {
    status: string;
    accessToken: string;
    data: {
        user: User;
    };
}

declare const authService: {
    login: (credentials: { username: string; password: string }) => Promise<AuthResponse>;
    logout: () => Promise<void>;
    getToken: () => string | null;
    getCurrentUser: () => User | null;
    isAuthenticated: () => boolean;
};

export default authService;''')
code3.style = 'No Spacing'
for run in code3.runs:
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 128)

doc.add_page_break()

# File 4
doc.add_heading('4. File Login_Parents.tsx - Giao Diện Login', 2)
doc.add_paragraph('📁 Đường dẫn: src/pages/parent/Login_Parents.tsx')
doc.add_paragraph('🎯 Chức năng:')

features = doc.add_paragraph()
features.style = 'List Bullet'
features.add_run('Hiển thị form đăng nhập (username + password)')
features = doc.add_paragraph()
features.style = 'List Bullet'
features.add_run('Gọi authService.login() khi submit')
features = doc.add_paragraph()
features.style = 'List Bullet'
features.add_run('Xử lý lỗi và hiển thị thông báo')
features = doc.add_paragraph()
features.style = 'List Bullet'
features.add_run('Navigate đến Dashboard khi thành công')

doc.add_paragraph('💻 Code chính:')
code4 = doc.add_paragraph('''import { useState } from 'react';
import { useNavigate } from 'react-router-dom';
import authService from '../../services/authService';

export default function Login() {
  const navigate = useNavigate();
  const [formData, setFormData] = useState({ username: '', password: '' });
  const [error, setError] = useState('');
  const [loading, setLoading] = useState(false);

  const handleSubmit = async (e) => {
    e.preventDefault();
    setError('');
    setLoading(true);

    try {
      // Gọi API login
      await authService.login(formData);
      
      // Chuyển đến Dashboard khi thành công
      navigate('/parent/dashboard');
    } catch (err) {
      console.error("Login Error:", err);
      
      // Xử lý các loại lỗi khác nhau
      let msg = 'Đăng nhập thất bại.';
      if (err.code === "ERR_NETWORK") {
        msg = 'Lỗi kết nối server.';
      } else if (err.response?.status === 401 || err.response?.status === 400) {
        msg = 'Sai thông tin đăng nhập.';
      } else if (err.response?.data?.message) {
        msg = err.response.data.message;
      }
      setError(msg);
    } finally {
      setLoading(false);
    }
  };

  return (
    <form onSubmit={handleSubmit}>
      <input
        type="text"
        value={formData.username}
        onChange={(e) => setFormData({ ...formData, username: e.target.value })}
        placeholder="Username"
      />
      <input
        type="password"
        value={formData.password}
        onChange={(e) => setFormData({ ...formData, password: e.target.value })}
        placeholder="Password"
      />
      <button type="submit" disabled={loading}>
        {loading ? 'Đang xử lý...' : 'Đăng nhập'}
      </button>
      {error && <p className="error">{error}</p>}
    </form>
  );
}''')
code4.style = 'No Spacing'
for run in code4.runs:
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0, 0, 128)

doc.add_page_break()

# Section 3: Flow
doc.add_heading('III. FLOW HOẠT ĐỘNG', 1)
doc.add_paragraph('Quy trình xác thực diễn ra như sau:')

step = doc.add_paragraph()
step.style = 'List Number'
step.add_run('User nhập username & password vào form Login_Parents.tsx')

step = doc.add_paragraph()
step.style = 'List Number'
step.add_run('Click button "Đăng nhập" → gọi handleSubmit()')

step = doc.add_paragraph()
step.style = 'List Number'
step.add_run('handleSubmit() gọi authService.login(formData)')

step = doc.add_paragraph()
step.style = 'List Number'
step.add_run('authService.login() sử dụng api.js gửi POST /auth/signin')

step = doc.add_paragraph()
step.style = 'List Number'
step.add_run('Backend xác thực và trả về { accessToken, data: { user } }')

step = doc.add_paragraph()
step.style = 'List Number'
step.add_run('authService lưu token vào localStorage.setItem("token", accessToken)')

step = doc.add_paragraph()
step.style = 'List Number'
step.add_run('Lưu user info vào localStorage.setItem("user", JSON.stringify(user))')

step = doc.add_paragraph()
step.style = 'List Number'
step.add_run('Navigate sang /parent/dashboard')

step = doc.add_paragraph()
step.style = 'List Number'
step.add_run('Các request tiếp theo tự động có token trong header (nhờ api interceptor)')

doc.add_page_break()

# Section 4: Format
doc.add_heading('IV. FORMAT DỮ LIỆU GỬI VÀ NHẬN', 1)

doc.add_heading('Request Format (Gửi lên backend):', 2)
req = doc.add_paragraph('''POST /auth/signin
Content-Type: application/json

{
  "username": "nguyenvanan@gmail.com",
  "password": "Parent@123"
}''')
req.style = 'No Spacing'
for run in req.runs:
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

doc.add_heading('Response Format (Nhận từ backend):', 2)
res = doc.add_paragraph('''{
  "status": "success",
  "accessToken": "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9...",
  "data": {
    "user": {
      "_id": "693301928d3adfd7b6cce7be",
      "name": "Nguyễn Văn An",
      "email": "nguyenvanan@gmail.com",
      "phoneNumber": "0901234001",
      "role": "Parent"
    }
  }
}''')
res.style = 'No Spacing'
for run in res.runs:
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

# Section 5: Testing
doc.add_heading('V. CÁCH KIỂM TRA', 1)

test = doc.add_paragraph()
test.style = 'List Number'
test.add_run('Mở ứng dụng và vào trang /parent/login')

test = doc.add_paragraph()
test.style = 'List Number'
test.add_run('Nhập username và password hợp lệ')

test = doc.add_paragraph()
test.style = 'List Number'
test.add_run('Mở DevTools (F12) → Tab Console để xem log')

test = doc.add_paragraph()
test.style = 'List Number'
test.add_run('Click "Đăng nhập"')

test = doc.add_paragraph()
test.style = 'List Number'
test.add_run('Kiểm tra Console có log "siuu {status, accessToken, data}"')

test = doc.add_paragraph()
test.style = 'List Number'
test.add_run('Mở DevTools → Tab Application → Local Storage → Kiểm tra key "token" và "user"')

test = doc.add_paragraph()
test.style = 'List Number'
test.add_run('Nếu thành công, trang tự động chuyển sang /parent/dashboard')

# Section 6: Common Errors
doc.add_heading('VI. LỖI THƯỜNG GẶP & CÁCH XỬ LÝ', 1)

doc.add_heading('1. Lỗi 401 - Unauthorized', 2)
doc.add_paragraph('❌ Nguyên nhân: Sai username hoặc password')
doc.add_paragraph('✅ Giải pháp: Kiểm tra lại thông tin đăng nhập')

doc.add_heading('2. ERR_NETWORK', 2)
doc.add_paragraph('❌ Nguyên nhân: Không kết nối được backend')
doc.add_paragraph('✅ Giải pháp:')
sol = doc.add_paragraph()
sol.style = 'List Bullet'
sol.add_run('Kiểm tra backend đang chạy')
sol = doc.add_paragraph()
sol.style = 'List Bullet'
sol.add_run('Kiểm tra VITE_API_URL trong file .env')
sol = doc.add_paragraph()
sol.style = 'List Bullet'
sol.add_run('Kiểm tra CORS settings ở backend')

doc.add_heading('3. TypeError: Cannot read property "username"', 2)
doc.add_paragraph('❌ Nguyên nhân: Mismatch giữa frontend và backend về tên field')
doc.add_paragraph('✅ Giải pháp: Đảm bảo backend nhận { username, password } chứ không phải { email, password }')

# Section 7: Environment
doc.add_heading('VII. CẤU HÌNH ENVIRONMENT', 1)
doc.add_paragraph('Tạo file .env trong thư mục frontend:')
env = doc.add_paragraph('''VITE_API_URL=https://smart-school-bus-api.onrender.com/api/v1
VITE_SOCKET_URL=https://smart-school-bus-api.onrender.com''')
env.style = 'No Spacing'
for run in env.runs:
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

# Footer
doc.add_paragraph()
doc.add_paragraph()
footer = doc.add_paragraph('--- HẾT ---')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer.runs[0]
footer_run.font.size = Pt(12)
footer_run.bold = True

# Save
doc.save('c:/Users/Acer/OneDrive/Desktop/CNPM_A/CNPM-/HUONG_DAN_API_AUTHENTICATION_LOGIN.docx')
print("✅ Đã tạo file HUONG_DAN_API_AUTHENTICATION_LOGIN.docx thành công!")
